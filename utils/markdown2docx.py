#!/usr/bin/env python3

# This file is a modified version of the markdown2docx.py file from the Markdown2docx project.
# The original project is at https://pypi.org/project/Markdown2docx/
# The original file is licensed under the MIT license.

import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
from bs4 import BeautifulSoup
from PIL import Image
from docx.oxml.shared import OxmlElement, qn
from html.parser import HTMLParser
from .preprocess_markdown2docx import PreprocessMarkdown2docx
from pathlib import Path


class HtmlListParser(HTMLParser):
    """Parser for handling HTML lists and converting them to Word document lists."""

    list_level = -1
    lists = ["List Bullet", "List Bullet 2", "List Bullet 3"]
    doc = None  # the .docx document object
    spacing = "    "  # used if we run out of bullet levels
    spare_list = "○  "

    def handle_starttag(self, tag, attrs):
        if tag in ["ol", "ul"]:
            self.list_level += 1

    def handle_endtag(self, tag):
        if tag in ["ol", "ul"]:
            self.list_level -= 1

    def handle_data(self, data):
        data = data.strip()
        if data:
            if self.list_level in range(len(self.lists)):
                self.doc.add_paragraph(data, self.lists[self.list_level])
            else:
                self.doc.add_paragraph(
                    "        " + self.spacing * self.list_level + self.spare_list + data
                )


def find_page_width(doc):
    """Calculate page width in inches from Word document."""
    return float(doc.sections[0].page_width / 914400)


def do_table_of_contents(document):
    """Insert Table of Contents Field.
    The user will be asked to update the field when the docx is opened.
    """
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r'TOC \o "1-3" \h \z \u'
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "separate")
    fld_char3 = OxmlElement("w:t")
    fld_char3.text = "Right-click to update field."
    fld_char2.append(fld_char3)
    fld_char4 = OxmlElement("w:fldChar")
    fld_char4.set(qn("w:fldCharType"), "end")
    r_element = run._r
    r_element.append(fld_char)
    r_element.append(instr_text)
    r_element.append(fld_char2)
    r_element.append(fld_char4)


def do_table(doc, table_in, style):
    """Draw a table in the Word document."""
    the_header = table_in.find("thead")
    the_column_names = the_header.find_all("th")
    the_data = table_in.find_all("td")
    n_cols = len(the_column_names)
    n_rows = int(len(the_data) / n_cols)
    this_table = doc.add_table(rows=n_rows + 1, cols=n_cols, style=style)

    # Add headers
    row = this_table.rows[0].cells
    for h_index, header in enumerate(the_column_names):
        row[h_index].text = "" if header.text == "" else header.string

    # Add data
    row_index = 0
    for d_index, data in enumerate(the_data):
        if not d_index % n_cols:
            row = this_table.rows[row_index + 1].cells
            row_index += 1
        row[d_index % n_cols].text = "" if data.text == "" else data.string


def find_image_size(image_file):
    """Get the dimensions of an image file."""
    return Image.open(image_file).size


def do_paragraph(
    line,
    doc,
    page_width_inches,
    style_body,
    assumed_pixels_per_inch=200,
    picture_fraction_of_width=0.7,
):
    """Process a paragraph, handling both text and images."""
    is_image = line.find("img")
    if is_image is not None:
        image_source = is_image["src"]
        w, h = find_image_size(image_source)
        w_in_inches = w / assumed_pixels_per_inch
        picture_width_inches = page_width_inches * picture_fraction_of_width
        chosen_width = min(picture_width_inches, w_in_inches)
        doc.add_picture(image_source, width=docx.shared.Inches(chosen_width))
        return
    doc.add_paragraph(line.text.strip(), style=style_body)


def do_pre_code(line, doc, style_quote_table):
    """Format preformatted code blocks."""
    table = doc.add_table(rows=1, cols=1, style=style_quote_table)
    cell = table.cell(0, 0)
    cell.text = line.text
    paragraphs = cell.paragraphs
    paragraph = paragraphs[0]
    run_obj = paragraph.runs
    run = run_obj[0]
    font = run.font
    font.size = Pt(10)
    font.name = "Courier"


def do_fake_horizontal_rule(doc, length_of_line=80, c="_"):
    """Add a horizontal rule to the document."""
    paragraph = doc.add_paragraph(c * length_of_line)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


class Markdown2docx:
    """Main class for converting markdown to docx."""

    def __init__(self, project_or_text, is_file=True, output_file=None):
        # Initialize document styles
        self.style_table = "Medium Shading 1 Accent 3"
        self.style_quote = "Body Text"
        self.style_body = "Body Text"
        self.style_quote_table = "Table Grid"
        self.toc_indicator = "contents"

        # Initialize file paths and content
        if is_file:
            self.project = project_or_text
            self.infile = ".".join([project_or_text, "md"])
            self.outfile = (
                ".".join([project_or_text, "docx"])
                if output_file is None
                else output_file
            )
            self.html_out_file = ".".join([project_or_text, "html"])
            self.markdown = self._read_markdown()
        else:
            self.project = None
            self.infile = None
            self.outfile = output_file if output_file else "output.docx"
            self.html_out_file = "output.html"
            self.markdown = project_or_text

        # Create document
        self.doc = docx.Document()
        self.page_width_inches = find_page_width(self.doc)

        # Convert to HTML and create soup
        self.html = markdown2.markdown(
            self.markdown,
            extras=["fenced-code-blocks", "code-friendly", "wiki-tables", "tables"],
        )
        self.soup = BeautifulSoup(self.html, "html.parser")

    def _read_markdown(self):
        """Read markdown from file."""
        with open(self.infile, "r", encoding="utf-8") as f:
            return f.read()

    def _process_soup(self):
        """Process the HTML soup to create the Word document."""
        list_of_tables = self.soup.find_all("table")
        table_counter = 0
        table_of_contents_done = 0

        for line in self.soup:
            if line is not None:
                # Handle table of contents
                if (
                    str(line).lower().find(self.toc_indicator) >= 0
                    and table_of_contents_done < 2
                ):
                    table_of_contents_done += 1
                    if table_of_contents_done == 2:
                        do_table_of_contents(self.doc)
                        continue

                # Handle emphasis
                if line.find("em"):
                    try:
                        self.doc.add_paragraph(line.text.strip(), style=self.style_body)
                    except AttributeError:
                        pass
                    continue

                # Handle horizontal rules
                if str(line).lower().find("<hr/>") == 0:
                    do_fake_horizontal_rule(self.doc)
                    continue

                # Handle headings
                if line.name == "h1":
                    self.doc.add_heading(line, 0)
                    continue
                if line.name == "h2":
                    self.doc.add_heading(line, 1)
                    continue
                if line.name == "h3":
                    self.doc.add_heading(line, 2)
                    continue
                if line.name == "h4":
                    self.doc.add_heading(line, 3)
                    continue

                # Handle paragraphs and images
                if line.name == "p":
                    do_paragraph(
                        line, self.doc, self.page_width_inches, self.style_body
                    )
                    continue

                # Handle code blocks
                if line.name == "pre":
                    do_pre_code(line, self.doc, self.style_quote_table)
                    continue

                # Handle tables
                if line.name == "table":
                    this_table_in = list_of_tables[table_counter]
                    do_table(self.doc, this_table_in, self.style_table)
                    table_counter += 1
                    continue

                # Handle lists
                if line.name in ["ul", "ol"]:
                    parser = HtmlListParser()
                    if line.name == "ol":
                        HtmlListParser.lists = [
                            "List Number",
                            "List Number 2",
                            "List Number 3",
                        ]
                        HtmlListParser.spare_list = "#  "
                    parser.doc = self.doc
                    parser.feed(str(line))
                    continue

    def convert(self):
        """Convert markdown to docx."""
        self._process_soup()
        return self.doc

    def write_html(self):
        """Write the intermediate HTML to a file."""
        with open(self.html_out_file, "w", encoding="utf-8") as f:
            f.write(self.html)

    def save(self):
        """Save the Word document."""
        self.doc.save(self.outfile)


def convert_markdown_to_docx(
    markdown_input: str, output_file_path: Path = None, input_is_file=True
) -> Path:
    """Convert markdown to docx.

    Args:
        markdown_input (str): Either a file path without extension (if is_file=True)
                            or markdown text content (if is_file=False)
        output_file_path (Path, optional): Path for the output docx file.
                                   If None, uses input name for files or 'output.docx' for text.
        is_file (bool): Whether markdown_input is a file path (True) or markdown text (False)

    Returns:
        str: Path to the generated docx file
    """
    if input_is_file:
        # Handle file-based input (original behavior)
        ppm2w = PreprocessMarkdown2docx(markdown_input)
        markdown = ppm2w.get_all_but_macros()
        markdown = ppm2w.do_substitute_tokens(markdown)
        markdown = ppm2w.do_execute_commands(markdown)
    else:
        # Handle direct text input
        markdown = markdown_input

    # Make sure output file path exists
    output_file_path.parent.mkdir(parents=True, exist_ok=True)

    # Convert to docx
    converter = Markdown2docx(
        markdown_input, is_file=input_is_file, output_file=str(output_file_path)
    )
    converter.convert()
    converter.save()
    return Path(converter.outfile)
